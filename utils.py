from sklearn.metrics import classification_report, cohen_kappa_score
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams
from io import StringIO
from docx import Document
import pandas as pd
import openai
import time
from gptutils_10 import get_chatgpt_response, num_tokens_from_string

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        output_string = StringIO()
        extract_text_to_fp(file, output_string, laparams=LAParams())
        text = output_string.getvalue()
    return text #.split('\f')  # Splitting by page

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph.text.strip())
    text = ' '.join(paragraphs)
    return text

def classify_and_tag(name,classify_prompt, tag_prompt, token_limit = 16000):
    print('testing prompts:', name)
    # load projects and abstracts
    df = pd.read_excel("data/projects_to_review_w_abstracts.xlsx", index_col=0)
    df = df.dropna(subset='abstract')

    # classify as public health or not
    print('classifying abstracts')
    for n, row in df.iterrows():
        if n % 100 == 0:
            print('row', n, 'of',df.shape[0])
        prompt = classify_prompt + "Title:" + row['Project Title'] + ' Abstract:' + row['abstract']
        if num_tokens_from_string(prompt) > token_limit:
            print('note: truncating prompt due to token limit')
            prompt = prompt[:token_limit * 4]
        response = get_chatgpt_response([prompt]).return_response()
        if 'Verdict:' in response and 'Justification:' in response:
            response2 = response.replace('Verdict:', '').replace(f'\n\n', '').split('Justification:')
            verdict = response2[0].strip()
            justification = response2[1].strip()
            df.loc[n, 'public_health_flag'] = verdict
            df.loc[n, 'flag_justification'] = justification
        else:
            df.loc[n, 'public_health_flag'] = "Bad response"
            df.loc[n, 'flag_justification'] = response
        time.sleep(1)

    filt_df = df.loc[df['public_health_flag'].str.lower().str.contains('yes')]
    areas = [
        "Education (e.g., safe and healthy schools)",
        "Workforce (e.g., income inequality)",
        "Health system reform and improvement",
        "Health services (quality and delivery)",
        "Health equity",
        "Violence prevention and interruption",
        "Housing",
        "Aging and disability",
        "Food security",
        "Maternal and child health and welfare",
        "Substance use disorder",
        "Mental health and suicide prevention",
        "Diseases and vaccinations",
        "Environmental justice (e.g. water quality, pollution)",
        "Sexual and reproductive health",
        "Community design/urban planning/transportation",
        "None"
    ]

    print('tagging abstracts')
    for n, row in filt_df.iterrows():
        if n % 100 == 0:
            print('row', n, 'of', filt_df.shape[0])
        prompt = tag_prompt + "Title:" + row['Project Title'] + ' Abstract:' + row['abstract']
        if num_tokens_from_string(prompt) > token_limit:
            print('note: truncating prompt due to token limit')
            prompt = prompt[:token_limit*4]
        response = get_chatgpt_response([prompt]).return_response()
        df.loc[n, 'public_health_cat'] = response
        time.sleep(1)

    def clean_cats(s):
        for a in areas:
            if not pd.isna(s) and a in s:
                clean_s = a
                return clean_s
        return 'None'

    df['public_health_cat'] = df['public_health_cat'].apply(clean_cats)

    # in the manual sample, a lot of false positives tend to be the education

    df.to_excel(f'data/projects_chatgpt_flags_tagged {name}.xlsx', index=False)

    print('analyzing performance')
    gpt_df = df.copy()
    gpt_df = gpt_df.loc[gpt_df.public_health_flag != 'Unknown']
    human_df = pd.read_excel("data/human_tags.xlsx")
    human_df = human_df.loc[human_df['Public Health?'].isna() == False]
    human_df = human_df[['Project Number', 'Public Health?', 'If Y, subcategory']]
    human_df = human_df.rename({'Public Health?': 'human_ph_flag',
                                'If Y, subcategory': 'human_ph_cat'}, axis=1)
    human_df['human_ph_flag'] = human_df['human_ph_flag'].replace({'N': 'No', 'Y': 'Yes'})
    gpt_df = gpt_df[['Project Number', 'abstract', 'public_health_flag', 'flag_justification', 'public_health_cat']]
    gpt_df.columns = ['Project Number', 'abstract', 'bot_ph_flag', 'flag_justification', 'bot_ph_cat']
    df = gpt_df.merge(human_df, on='Project Number')
    df['cat_match'] = df.human_ph_cat == df.bot_ph_cat
    df.loc[(df.bot_ph_flag != 'Yes') | (df.human_ph_flag != 'Ye'), 'cat_match'] = pd.NA

    result_df = pd.DataFrame()
    # print('human-bot public health flag cohen K')
    result_df.loc['human-bot public health flag', 'cohen K'] = cohen_kappa_score(df['human_ph_flag'], df['bot_ph_flag'])
    result_df.loc['human-bot public health flag', 'N'] = df.shape[0]

    h2_df = pd.read_excel('data/for_second_tagging.xlsx')
    h2_df = h2_df[['Project Number', 'abstract', 'Public Health?', 'If Y, subcategory', 'ph_flag_2', 'ph_cat2']]
    h2_df = h2_df.rename({'Public Health?': 'human_ph_flag',
                          'If Y, subcategory': 'human_ph_cat'}, axis=1)
    # print('human-human public health flag cohen K')
    result_df.loc['human-human public health flag', 'cohen K'] = cohen_kappa_score(h2_df['human_ph_flag'],
                                                                                   h2_df['ph_flag_2'])
    result_df.loc['human-human public health flag', 'N'] = h2_df.shape[0]

    filt_df = df.loc[(df.bot_ph_flag == 'Yes') & (df.human_ph_flag == 'Yes') & (df['human_ph_cat'].isna() == False) & (
                df['bot_ph_cat'].isna() == False)]
    # print('human-bot public health category cohen K')
    result_df.loc['human-bot public health category', 'cohen K'] = cohen_kappa_score(filt_df['human_ph_cat'],
                                                                                     filt_df['bot_ph_cat'])
    result_df.loc['human-bot public health category', 'N'] = filt_df.shape[0]

    filt2_df = h2_df.loc[(h2_df.ph_flag_2 == 'Y') & (h2_df.human_ph_flag == 'Y')]
    # print('human-human public health category cohen K')
    result_df.loc['human-human public health category', 'cohen K'] = cohen_kappa_score(filt2_df['human_ph_cat'],
                                                                                       filt2_df['ph_cat2'])
    result_df.loc['human-human public health category', 'N'] = filt2_df.shape[0]
    result_df.to_excel(f'output/performance test {name}.xlsx')

    pd.crosstab(df.bot_ph_flag, df.human_ph_flag).to_csv(f'output/disagreement xtab {name}.csv')

    flag_diffs = df.loc[(df.bot_ph_flag != df.human_ph_flag)]
    flag_diffs.to_excel(f'output/flag disagreements {name}.xlsx', index=False)